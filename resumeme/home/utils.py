import os
from django.conf import settings
from django.core.mail import send_mail
from django.template.loader import render_to_string
from django.utils.html import strip_tags
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import os
from django.conf import settings
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from django.conf import settings
import openai
from django.conf import settings
from typing import List, Dict, Any
import json


class AIHelper:
    def __init__(self):
        openai.api_key = settings.OPENAI_API_KEY
        self.model = "gpt-3.5-turbo"
        self.max_tokens = 1000
        self.temperature = 0.7

    def improve_text(self, text: str, context: str = "") -> str:
        """
        Improve the given text while maintaining its core meaning
        """
        prompt = f"""
        Please improve this text while maintaining its core meaning and professional tone.
        Make it more impactful and engaging, but keep it concise.

        Context: {context}

        Original text:
        {text}
        """

        try:
            response = openai.ChatCompletion.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are a professional resume writer and career coach."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=self.max_tokens,
                temperature=self.temperature
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            print(f"Error improving text: {e}")
            return text

    def generate_bullet_points(self, description: str, count: int = 3) -> List[str]:
        """
        Generate bullet points from a description
        """
        prompt = f"""
        Convert this description into {count} impactful bullet points for a resume.
        Focus on achievements and quantifiable results where possible.

        Description:
        {description}
        """

        try:
            response = openai.ChatCompletion.create(
                model=self.model,
                messages=[
                    {"role": "system",
                     "content": "You are a professional resume writer specializing in achievement-based bullet points."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=self.max_tokens,
                temperature=self.temperature
            )

            # Process the response into a list of bullet points
            bullets = response.choices[0].message.content.strip().split('\n')
            # Clean up bullets (remove leading dashes or bullets)
            bullets = [b.lstrip('•- ').strip() for b in bullets if b.strip()]
            return bullets[:count]
        except Exception as e:
            print(f"Error generating bullet points: {e}")
            return []

    def suggest_skills(self, job_description: str, count: int = 5) -> List[str]:
        """
        Suggest relevant skills based on a job description
        """
        prompt = f"""
        Based on this job description, suggest {count} relevant skills that should be included in a resume.
        Format the response as a simple list of skills.

        Job Description:
        {job_description}
        """

        try:
            response = openai.ChatCompletion.create(
                model=self.model,
                messages=[
                    {"role": "system",
                     "content": "You are a career coach specializing in skill identification and matching."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=self.max_tokens,
                temperature=self.temperature
            )

            # Process the response into a list of skills
            skills = response.choices[0].message.content.strip().split('\n')
            # Clean up skills (remove leading dashes or bullets)
            skills = [s.lstrip('•- ').strip() for s in skills if s.strip()]
            return skills[:count]
        except Exception as e:
            print(f"Error suggesting skills: {e}")
            return []

    def analyze_resume(self, resume_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Analyze a resume and provide feedback and suggestions
        """
        prompt = f"""
        Please analyze this resume and provide feedback in the following areas:
        1. Overall Impact
        2. Areas for Improvement
        3. Suggested Enhancements
        4. Keywords Analysis

        Resume Data:
        {json.dumps(resume_data, indent=2)}

        Provide the analysis in JSON format with these categories.
        """

        try:
            response = openai.ChatCompletion.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are an expert resume analyst and career coach."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=self.max_tokens,
                temperature=self.temperature
            )

            # Parse the JSON response
            analysis = json.loads(response.choices[0].message.content)
            return analysis
        except Exception as e:
            print(f"Error analyzing resume: {e}")
            return {
                "overall_impact": "Analysis unavailable",
                "areas_for_improvement": [],
                "suggested_enhancements": [],
                "keywords_analysis": []
            }

    def get_job_title_suggestions(self, skills: List[str], experience_years: int) -> List[str]:
        """
        Suggest job titles based on skills and experience
        """
        prompt = f"""
        Based on these skills and years of experience, suggest 5 relevant job titles.

        Skills: {', '.join(skills)}
        Years of Experience: {experience_years}

        Provide just the job titles, one per line.
        """

        try:
            response = openai.ChatCompletion.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are a career counselor specializing in job matching."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=self.max_tokens,
                temperature=self.temperature
            )

            # Process the response into a list of job titles
            titles = response.choices[0].message.content.strip().split('\n')
            return [t.strip() for t in titles if t.strip()]
        except Exception as e:
            print(f"Error getting job title suggestions: {e}")
            return []

    def get_achievement_suggestions(self, role: str, responsibilities: str) -> List[str]:
        """
        Suggest potential achievements based on job role and responsibilities
        """
        prompt = f"""
        Based on this job role and responsibilities, suggest 3 potential achievements or accomplishments
        that could be highlighted on a resume. Focus on quantifiable results.

        Role: {role}
        Responsibilities: {responsibilities}
        """

        try:
            response = openai.ChatCompletion.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are a resume writer specializing in achievement-based writing."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=self.max_tokens,
                temperature=self.temperature
            )

            # Process the response into a list of achievements
            achievements = response.choices[0].message.content.strip().split('\n')
            return [a.lstrip('•- ').strip() for a in achievements if a.strip()]
        except Exception as e:
            print(f"Error getting achievement suggestions: {e}")
            return []


class WordExporter:
    def __init__(self, resume):
        self.resume = resume
        self.document = Document()
        self.output_dir = os.path.join(settings.MEDIA_ROOT, 'docs')
        os.makedirs(self.output_dir, exist_ok=True)

    def generate(self):
        """Generate Word document from resume data"""
        # Set up document margins
        sections = self.document.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Add header/personal info
        self.add_header()

        # Add sections
        for section in self.resume.sections.all().order_by('order'):
            self.format_section(section)

        # Save the document
        filename = f'resume_{self.resume.id}.docx'
        filepath = os.path.join(self.output_dir, filename)
        self.document.save(filepath)
        return os.path.join('docs', filename)

    def add_header(self):
        """Add personal information section"""
        personal_info = self.get_personal_info()
        if personal_info:
            # Name
            if 'name' in personal_info:
                name = self.document.add_paragraph()
                name_run = name.add_run(personal_info['name'])
                name_run.font.size = Pt(24)
                name_run.font.bold = True
                name.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Title
            if 'title' in personal_info:
                title = self.document.add_paragraph()
                title_run = title.add_run(personal_info['title'])
                title_run.font.size = Pt(16)
                title_run.font.italic = True
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Contact Info
            contact_info = []
            if 'email' in personal_info:
                contact_info.append(personal_info['email'])
            if 'phone' in personal_info:
                contact_info.append(personal_info['phone'])
            if 'location' in personal_info:
                contact_info.append(personal_info['location'])

            if contact_info:
                contact = self.document.add_paragraph()
                contact.add_run(' | '.join(contact_info))
                contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

            self.document.add_paragraph()  # Add spacing

    def format_section(self, section):
        """Format a resume section"""
        # Add section title
        title = self.document.add_paragraph()
        title_run = title.add_run(section.title)
        title_run.font.size = Pt(14)
        title_run.font.bold = True

        # Add horizontal line
        title.paragraph_format.bottom_border.width = Pt(1)

        # Format content based on section type
        if section.type == 'experience':
            self.format_experience(section)
        elif section.type == 'education':
            self.format_education(section)
        elif section.type == 'skills':
            self.format_skills(section)
        elif section.type == 'languages':
            self.format_languages(section)

        self.document.add_paragraph()  # Add spacing between sections

    def format_experience(self, section):
        """Format work experience section"""
        for exp in section.experiences.all():
            # Company and position
            p = self.document.add_paragraph()
            company = p.add_run(f"{exp.company}")
            company.font.bold = True
            p.add_run(" - ")
            p.add_run(exp.position)

            # Dates
            date_text = f"{exp.start_date.strftime('%B %Y')} - "
            date_text += "Present" if exp.is_current else exp.end_date.strftime('%B %Y')
            dates = self.document.add_paragraph()
            dates.add_run(date_text).italic = True

            # Description
            if exp.description:
                desc = self.document.add_paragraph()
                desc.add_run(exp.description)

            self.document.add_paragraph()  # Add spacing

    def format_education(self, section):
        """Format education section"""
        for edu in section.education.all():
            p = self.document.add_paragraph()
            institution = p.add_run(f"{edu.institution}")
            institution.font.bold = True
            p.add_run(f" - {edu.degree}")
            self.document.add_paragraph()  # Add spacing

    def format_skills(self, section):
        """Format skills section"""
        skills = section.skills.all()
        if skills:
            p = self.document.add_paragraph()
            skill_text = ", ".join([
                f"{skill.name} ({skill.level})"
                for skill in skills
            ])
            p.add_run(skill_text)

    def format_languages(self, section):
        """Format languages section"""
        languages = section.languages.all()
        if languages:
            p = self.document.add_paragraph()
            lang_text = ", ".join([
                f"{lang.name} ({lang.level})"
                for lang in languages
            ])
            p.add_run(lang_text)

    def get_personal_info(self):
        """Get personal information from resume"""
        personal_section = self.resume.sections.filter(type='personal').first()
        if personal_section:
            return personal_section.content
        return {}

    def save_to_file(self, filename=None):
        """Save Word document to a file"""
        return self.generate()


class PDFExporter:
    def __init__(self, resume):
        self.resume = resume
        self.styles = getSampleStyleSheet()
        self.output_dir = os.path.join(settings.MEDIA_ROOT, 'pdfs')
        os.makedirs(self.output_dir, exist_ok=True)

        # Custom styles
        self.styles.add(ParagraphStyle(
            name='SectionTitle',
            parent=self.styles['Heading1'],
            fontSize=14,
            spaceAfter=10,
            textColor=colors.HexColor('#2c3e50')
        ))

        self.styles.add(ParagraphStyle(
            name='SubTitle',
            parent=self.styles['Heading2'],
            fontSize=12,
            textColor=colors.HexColor('#3498db')
        ))

    def generate(self):
        """Generate PDF from resume data"""
        filename = f'resume_{self.resume.id}.pdf'
        filepath = os.path.join(self.output_dir, filename)

        # Create the PDF document
        doc = SimpleDocTemplate(
            filepath,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )

        # Create the story (content) for the PDF
        story = []

        # Add header/personal info
        personal_info = self.get_personal_info()
        if personal_info:
            story.extend(self.create_header(personal_info))
            story.append(Spacer(1, 20))

        # Add sections
        for section in self.resume.sections.all().order_by('order'):
            story.extend(self.format_section(section))
            story.append(Spacer(1, 20))

        # Build the PDF
        doc.build(story)
        return os.path.join('pdfs', filename)

    def get_personal_info(self):
        """Get personal information from resume"""
        personal_section = self.resume.sections.filter(type='personal').first()
        if personal_section:
            return personal_section.content
        return {}

    def create_header(self, personal_info):
        """Create header section with personal information"""
        elements = []

        # Name
        if 'name' in personal_info:
            elements.append(Paragraph(
                personal_info['name'],
                ParagraphStyle(
                    'Name',
                    parent=self.styles['Title'],
                    fontSize=24,
                    alignment=1  # Center alignment
                )
            ))

        # Title/Position
        if 'title' in personal_info:
            elements.append(Paragraph(
                personal_info['title'],
                ParagraphStyle(
                    'Title',
                    parent=self.styles['SubTitle'],
                    fontSize=16,
                    alignment=1
                )
            ))

        # Contact info
        contact_info = []
        if 'email' in personal_info:
            contact_info.append(personal_info['email'])
        if 'phone' in personal_info:
            contact_info.append(personal_info['phone'])
        if 'location' in personal_info:
            contact_info.append(personal_info['location'])

        if contact_info:
            elements.append(Paragraph(
                ' | '.join(contact_info),
                ParagraphStyle(
                    'ContactInfo',
                    parent=self.styles['Normal'],
                    alignment=1,
                    spaceAfter=20
                )
            ))

        return elements

    def format_section(self, section):
        """Format a resume section"""
        elements = []

        # Section title
        elements.append(Paragraph(section.title, self.styles['SectionTitle']))

        # Section content based on type
        if section.type == 'experience':
            elements.extend(self.format_experience(section))
        elif section.type == 'education':
            elements.extend(self.format_education(section))
        elif section.type == 'skills':
            elements.extend(self.format_skills(section))
        elif section.type == 'languages':
            elements.extend(self.format_languages(section))

        return elements

    def format_experience(self, section):
        """Format work experience section"""
        elements = []
        for exp in section.experiences.all():
            # Company and position
            elements.append(Paragraph(
                f"<b>{exp.company}</b> - {exp.position}",
                self.styles['SubTitle']
            ))

            # Dates
            date_text = f"{exp.start_date.strftime('%B %Y')} - "
            date_text += "Present" if exp.is_current else exp.end_date.strftime('%B %Y')
            elements.append(Paragraph(
                date_text,
                self.styles['Italic']
            ))

            # Description
            if exp.description:
                elements.append(Paragraph(
                    exp.description,
                    self.styles['Normal']
                ))

            elements.append(Spacer(1, 12))
        return elements

    def format_education(self, section):
        """Format education section"""
        elements = []
        for edu in section.education.all():
            elements.append(Paragraph(
                f"<b>{edu.institution}</b> - {edu.degree}",
                self.styles['SubTitle']
            ))
            elements.append(Spacer(1, 12))
        return elements

    def format_skills(self, section):
        """Format skills section"""
        elements = []
        skills = section.skills.all()
        if skills:
            skill_text = ", ".join([
                f"{skill.name} ({skill.level})"
                for skill in skills
            ])
            elements.append(Paragraph(
                skill_text,
                self.styles['Normal']
            ))
        return elements

    def format_languages(self, section):
        """Format languages section"""
        elements = []
        languages = section.languages.all()
        if languages:
            lang_text = ", ".join([
                f"{lang.name} ({lang.level})"
                for lang in languages
            ])
            elements.append(Paragraph(
                lang_text,
                self.styles['Normal']
            ))
        return elements

    def save_to_file(self, filename=None):
        """Save PDF to a file"""
        return self.generate()


def send_email_template(template_name, context, subject, recipient_list, from_email=None):
    """
    Send an email using a template
    """
    if from_email is None:
        from_email = settings.DEFAULT_FROM_EMAIL

    # Render HTML content
    html_content = render_to_string(template_name, context)

    # Create plain text content
    text_content = strip_tags(html_content)

    # Send email
    send_mail(
        subject=subject,
        message=text_content,
        from_email=from_email,
        recipient_list=recipient_list,
        html_message=html_content,
        fail_silently=False,
    )


def handle_uploaded_file(file, path):
    """
    Handle file upload and return the file path
    """
    filename = file.name
    filepath = os.path.join(path, filename)

    # Ensure directory exists
    os.makedirs(os.path.dirname(filepath), exist_ok=True)

    # Save file
    with open(filepath, 'wb+') as destination:
        for chunk in file.chunks():
            destination.write(chunk)

    return filepath


def generate_unique_slug(model_class, title, slug_field='slug'):
    """
    Generate a unique slug for a model instance
    """
    from django.utils.text import slugify

    slug = slugify(title)
    unique_slug = slug
    counter = 1

    while model_class.objects.filter(**{slug_field: unique_slug}).exists():
        unique_slug = f"{slug}-{counter}"
        counter += 1

    return unique_slug


def format_phone_number(phone):
    """
    Format phone number to a standard format
    """
    # Remove all non-numeric characters
    cleaned = ''.join(filter(str.isdigit, str(phone)))

    # Format based on length
    if len(cleaned) == 10:  # US number
        return f"({cleaned[:3]}) {cleaned[3:6]}-{cleaned[6:]}"
    elif len(cleaned) == 11 and cleaned[0] == '1':  # US number with country code
        return f"+1 ({cleaned[1:4]}) {cleaned[4:7]}-{cleaned[7:]}"
    else:
        return phone  # Return original if can't format


def get_file_extension(filename):
    """
    Get file extension from filename
    """
    return os.path.splitext(filename)[1].lower()


def is_valid_file_extension(filename, allowed_extensions):
    """
    Check if file extension is allowed
    """
    return get_file_extension(filename) in allowed_extensions


def get_client_ip(request):
    """
    Get client IP address from request
    """
    x_forwarded_for = request.META.get('HTTP_X_FORWARDED_FOR')
    if x_forwarded_for:
        ip = x_forwarded_for.split(',')[0]
    else:
        ip = request.META.get('REMOTE_ADDR')
    return ip


def truncate_string(value, max_length=50, suffix='...'):
    """
    Truncate a string to a maximum length
    """
    if len(value) <= max_length:
        return value
    truncated = value[:max_length - len(suffix)]
    return truncated.rstrip() + suffix


def format_file_size(size):
    """
    Format file size in bytes to human readable format
    """
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size < 1024:
            return f"{size:.1f} {unit}"
        size /= 1024
    return f"{size:.1f} TB"


def is_ajax(request):
    """
    Check if request is AJAX
    """
    return request.headers.get('X-Requested-With') == 'XMLHttpRequest'


class FileValidator:
    """
    Validate file uploads
    """

    def __init__(self, max_size=5 * 1024 * 1024, allowed_extensions=None):
        self.max_size = max_size
        self.allowed_extensions = allowed_extensions or ['.jpg', '.jpeg', '.png', '.pdf', '.doc', '.docx']

    def validate(self, file):
        if file.size > self.max_size:
            raise ValueError(f"File size cannot exceed {format_file_size(self.max_size)}")

        if not is_valid_file_extension(file.name, self.allowed_extensions):
            raise ValueError(f"File type not allowed. Allowed types: {', '.join(self.allowed_extensions)}")


def get_page_range(paginator, current_page, show_first_last=True):
    """
    Get page range for pagination with ellipsis
    """
    page_range = []

    if paginator.num_pages <= 5:
        page_range = paginator.page_range
    else:
        if current_page <= 3:
            page_range = range(1, 6)
            if show_first_last:
                page_range = list(page_range) + ['...', paginator.num_pages]
        elif current_page >= paginator.num_pages - 2:
            page_range = range(paginator.num_pages - 4, paginator.num_pages + 1)
            if show_first_last:
                page_range = [1, '...'] + list(page_range)
        else:
            page_range = range(current_page - 2, current_page + 3)
            if show_first_last:
                page_range = [1, '...'] + list(page_range) + ['...', paginator.num_pages]

    return page_range