from django.core.paginator import Paginator
from django.views.generic import ListView, DetailView
from django.db.models import Count, Q
from django.urls import reverse
from .models import ResumeTemplate, TemplateCategory, TemplateColor, UserTemplateSelection
from .forms import TemplateFilterForm
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from .models import ResumeTemplate, UserTemplateSelection
from resume.models import  Resume
from django.http import JsonResponse, Http404
from django.views.decorators.http import require_POST
import json
from home.models import Resume  # Assuming you have a Resume model
from django.shortcuts import render, get_object_or_404
from django.http import HttpResponse
import io
from django.template.loader import get_template
from xhtml2pdf import pisa  # You'll need to install this: pip install xhtml2pdf
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.views.generic import ListView, DetailView, CreateView, UpdateView
from django.contrib.auth.mixins import LoginRequiredMixin
from django.urls import reverse_lazy, reverse
from django.http import JsonResponse, HttpResponse
from django.template.loader import render_to_string
from django.utils.text import slugify
from django.contrib import messages
import json
from dashboard.models import ResumeUploadForm
# views.py
class AllTemplatesView(ListView):
    """View for browsing all resume templates"""
    model = ResumeTemplate
    template_name = 'templates_app/all_templates.html'
    context_object_name = 'templates'
    paginate_by = 12

    def get_queryset(self):
        queryset = super().get_queryset().filter(is_active=True)

        # Filter by category if provided
        category = self.request.GET.get('category')
        if category:
            queryset = queryset.filter(category=category)

        # Filter by style if provided
        style = self.request.GET.get('style')
        if style:
            queryset = queryset.filter(style=style)

        # Filter by color if provided
        color = self.request.GET.get('color')
        if color:
            queryset = queryset.filter(color=color)

        # Search by name or description
        search = self.request.GET.get('search')
        if search:
            queryset = queryset.filter(
                Q(name__icontains=search) |
                Q(description__icontains=search)
            )

        # Sort by popularity, newest, or name
        sort = self.request.GET.get('sort', 'popular')
        if sort == 'newest':
            queryset = queryset.order_by('-created_at')
        elif sort == 'name':
            queryset = queryset.order_by('name')
        else:  # Default to popular
            queryset = queryset.order_by('-popularity')

        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        # Add filter options to context
        context['categories'] = ResumeTemplate.objects.values_list('category', flat=True).distinct()
        context['styles'] = ResumeTemplate.objects.values_list('style', flat=True).distinct()
        context['colors'] = ResumeTemplate.objects.values_list('color', flat=True).distinct()

        # Add current filter selections
        context['selected_category'] = self.request.GET.get('category', '')
        context['selected_style'] = self.request.GET.get('style', '')
        context['selected_color'] = self.request.GET.get('color', '')
        context['selected_sort'] = self.request.GET.get('sort', 'popular')
        context['search_query'] = self.request.GET.get('search', '')

        return context


class SimpleTemplatesView(ListView):
    """View for browsing simple resume templates"""
    model = ResumeTemplate
    template_name = 'templates_app/simple_templates/simple_template.html'
    context_object_name = 'templates'

    def get_queryset(self):
        queryset = super().get_queryset().filter(category='simple')

        # Filter by style if provided
        style = self.request.GET.get('style')
        if style:
            queryset = queryset.filter(style=style)

        # Filter by color if provided
        color = self.request.GET.get('color')
        if color:
            queryset = queryset.filter(color=color)

        # Sort by popularity, newest, or name
        sort = self.request.GET.get('sort', 'popular')
        if sort == 'newest':
            queryset = queryset.order_by('-created_at')
        elif sort == 'name':
            queryset = queryset.order_by('name')
        else:  # Default to popular
            queryset = queryset.order_by('-popularity')

        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        # Add filter options to context
        context['styles'] = ResumeTemplate.objects.filter(category='simple').values_list('style', flat=True).distinct()
        context['colors'] = ResumeTemplate.objects.filter(category='simple').values_list('color', flat=True).distinct()

        # Add current filter selections
        context['selected_style'] = self.request.GET.get('style', '')
        context['selected_color'] = self.request.GET.get('color', '')
        context['selected_sort'] = self.request.GET.get('sort', 'popular')

        # Add category information
        context['category'] = 'simple'
        context['category_title'] = 'Simple Resume Templates'
        context['category_description'] = 'Clean and straightforward designs that focus on content.'

        return context


class ModernTemplatesView(ListView):
    """View for browsing modern resume templates"""
    model = ResumeTemplate
    template_name = 'templates_app/modern_templates/modern_template.html'
    context_object_name = 'templates'

    def get_queryset(self):
        queryset = super().get_queryset().filter(category='modern')

        # Filter by style if provided
        style = self.request.GET.get('style')
        if style:
            queryset = queryset.filter(style=style)

        # Filter by color if provided
        color = self.request.GET.get('color')
        if color:
            queryset = queryset.filter(color=color)

        # Sort by popularity, newest, or name
        sort = self.request.GET.get('sort', 'popular')
        if sort == 'newest':
            queryset = queryset.order_by('-created_at')
        elif sort == 'name':
            queryset = queryset.order_by('name')
        else:  # Default to popular
            queryset = queryset.order_by('-popularity')

        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        # Add filter options to context
        context['styles'] = ResumeTemplate.objects.filter(category='modern').values_list('style', flat=True).distinct()
        context['colors'] = ResumeTemplate.objects.filter(category='modern').values_list('color', flat=True).distinct()

        # Add current filter selections
        context['selected_style'] = self.request.GET.get('style', '')
        context['selected_color'] = self.request.GET.get('color', '')
        context['selected_sort'] = self.request.GET.get('sort', 'popular')

        # Add category information
        context['category'] = 'modern'
        context['category_title'] = 'Modern Resume Templates'
        context['category_description'] = 'Contemporary designs with creative layouts and styling.'

        return context


class CreativeTemplatesView(ListView):
    """View for browsing creative resume templates"""
    model = ResumeTemplate
    template_name = 'templates_app/creative_templates/creative_template.html'
    context_object_name = 'templates'

    def get_queryset(self):
        queryset = super().get_queryset().filter(category='creative')

        # Filter by style if provided
        style = self.request.GET.get('style')
        if style:
            queryset = queryset.filter(style=style)

        # Filter by color if provided
        color = self.request.GET.get('color')
        if color:
            queryset = queryset.filter(color=color)

        # Sort by popularity, newest, or name
        sort = self.request.GET.get('sort', 'popular')
        if sort == 'newest':
            queryset = queryset.order_by('-created_at')
        elif sort == 'name':
            queryset = queryset.order_by('name')
        else:  # Default to popular
            queryset = queryset.order_by('-popularity')

        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        # Add filter options to context
        context['styles'] = ResumeTemplate.objects.filter(category='creative').values_list('style',
                                                                                           flat=True).distinct()
        context['colors'] = ResumeTemplate.objects.filter(category='creative').values_list('color',
                                                                                           flat=True).distinct()

        # Add current filter selections
        context['selected_style'] = self.request.GET.get('style', '')
        context['selected_color'] = self.request.GET.get('color', '')
        context['selected_sort'] = self.request.GET.get('sort', 'popular')

        # Add category information
        context['category'] = 'creative'
        context['category_title'] = 'Creative Resume Templates'
        context['category_description'] = 'Bold and unique designs that help you stand out.'

        return context

def template_detail(request, template_slug):
    template = get_object_or_404(ResumeTemplate, slug=template_slug, is_active=True)

    # Get sample data for preview
    sample_data = {
        'personal_info': {
            'full_name': 'John Doe',
            'job_title': 'Software Engineer',
            'email': 'john.doe@example.com',
            'phone': '(123) 456-7890',
            'address': 'New York, NY',
            'summary': 'Experienced software engineer with a passion for developing innovative solutions...'
        },
        'experiences': [
            {
                'company': 'Tech Solutions Inc.',
                'position': 'Senior Software Engineer',
                'location': 'New York, NY',
                'start_date': '2018-01-01',
                'end_date': None,
                'current': True,
                'description': 'Led development of cloud-based applications...'
            },
            {
                'company': 'Digital Innovations',
                'position': 'Software Developer',
                'location': 'Boston, MA',
                'start_date': '2015-03-01',
                'end_date': '2017-12-31',
                'current': False,
                'description': 'Developed and maintained web applications...'
            }
        ],
        'educations': [
            {
                'institution': 'MIT',
                'degree': 'Master of Science',
                'field_of_study': 'Computer Science',
                'location': 'Cambridge, MA',
                'start_date': '2013-09-01',
                'end_date': '2015-05-31',
                'current': False,
                'description': ''
            }
        ],
        'skills': [
            {'name': 'JavaScript', 'level': 5},
            {'name': 'Python', 'level': 4},
            {'name': 'React', 'level': 4},
            {'name': 'Node.js', 'level': 3},
            {'name': 'SQL', 'level': 4}
        ],
        'languages': [
            {'name': 'English', 'proficiency': 'native'},
            {'name': 'Spanish', 'proficiency': 'professional'}
        ],
        'projects': [
            {
                'title': 'E-commerce Platform',
                'description': 'Developed a full-stack e-commerce platform...',
                'url': 'https://github.com/johndoe/ecommerce',
                'start_date': '2019-06-01',
                'end_date': '2019-12-31'
            }
        ],
        'certifications': [
            {
                'name': 'AWS Certified Solutions Architect',
                'issuing_organization': 'Amazon Web Services',
                'issue_date': '2020-01-15',
                'expiration_date': '2023-01-15',
                'credential_id': 'AWS-123456',
                'credential_url': 'https://aws.amazon.com/certification/verify'
            }
        ]
    }
    try:
        template = get_object_or_404(ResumeTemplate, slug=template_slug)
        return render(request, 'templates_app/template_detail.html', {
            'template': template,
        })
    except Http404:
        # Get all available templates to suggest alternatives
        available_templates = ResumeTemplate.objects.all()[:5]  # Get first 5 templates
        return render(request, 'templates_app/template_not_found.html', {
            'requested_slug': template_slug,
            'available_templates': available_templates,
        })

    # Get related templates (same category)
    related_templates = ResumeTemplate.objects.filter(
        category=template.category,
        is_active=True
    ).exclude(id=template.id).order_by('?')[:3]

    # Get color schemes for customization preview
    color_schemes = [
        {'name': 'Default', 'primary': '#3b82f6', 'secondary': '#6c757d', 'accent': '#10b981'},
        {'name': 'Professional', 'primary': '#1e40af', 'secondary': '#475569', 'accent': '#0369a1'},
        {'name': 'Creative', 'primary': '#7c3aed', 'secondary': '#6b7280', 'accent': '#ec4899'},
        {'name': 'Modern', 'primary': '#0f766e', 'secondary': '#4b5563', 'accent': '#f59e0b'},
        {'name': 'Classic', 'primary': '#b91c1c', 'secondary': '#374151', 'accent': '#ca8a04'}
    ]

    return render(request, 'templates_app/latest_template_detail.html', {
        'template': template,
        'sample_data': sample_data,
        'related_templates': related_templates,
        'color_schemes': color_schemes
    })

def compare_templates(request):
    template_ids = request.GET.getlist('templates')

    if not template_ids or len(template_ids) < 2:
        # Default to comparing the first template from each category
        templates = []
        for category in TemplateCategory.objects.all():
            template = ResumeTemplate.objects.filter(category=category, is_active=True).first()
            if template:
                templates.append(template)

        if len(templates) < 2:
            # If still not enough, just get any two templates
            templates = ResumeTemplate.objects.filter(is_active=True)[:2]
    else:
        templates = ResumeTemplate.objects.filter(id__in=template_ids, is_active=True)

    # Get all templates for selection
    all_templates = ResumeTemplate.objects.filter(is_active=True).order_by('category__name', 'name')

    # Sample data for preview
    sample_data = {
        'personal_info': {
            'full_name': 'John Doe',
            'job_title': 'Software Engineer',
            'email': 'john.doe@example.com',
            'phone': '(123) 456-7890',
        }
    }

    return render(request, 'resume_builder/compare_templates.html', {
        'templates': templates,
        'all_templates': all_templates,
        'sample_data': sample_data
    })


def category_showcase(request, category_slug):
    category = get_object_or_404(TemplateCategory, slug=category_slug)
    templates = ResumeTemplate.objects.filter(category=category, is_active=True)

    return render(request, 'templates_app/category_showcase.html', {
        'category': category,
        'templates': templates
    })

def template_api(request):
    page = request.GET.get('page', 1)
    category = request.GET.get('category')
    search = request.GET.get('search')

    templates = ResumeTemplate.objects.filter(is_active=True)

    if category:
        templates = templates.filter(category__slug=category)

    if search:
        templates = templates.filter(
            Q(name__icontains=search) |
            Q(description__icontains=search) |
            Q(category__name__icontains=search)
        )

    paginator = Paginator(templates, 9)  # 9 templates per page
    page_obj = paginator.get_page(page)

    templates_data = []
    for template in page_obj:
        templates_data.append({
            'id': template.id,
            'name': template.name,
            'slug': template.slug,
            'description': template.description[:100] + '...' if len(
                template.description) > 100 else template.description,
            'preview_image': template.preview_image.url,
            'category_name': template.category.name,
            'category_slug': template.category.slug,
            'is_premium': template.is_premium
        })

    return JsonResponse({
        'templates': templates_data,
        'has_next': page_obj.has_next(),
        'total_pages': paginator.num_pages
    })


def template_detail(request, template_slug):
    template = get_object_or_404(ResumeTemplate, slug=template_slug, is_active=True)

    # Get sample data for preview
    sample_data = {
        'personal_info': {
            'full_name': 'John Doe',
            'job_title': 'Software Engineer',
            'email': 'john.doe@example.com',
            'phone': '(123) 456-7890',
            'address': 'New York, NY',
            'summary': 'Experienced software engineer with a passion for developing innovative solutions...'
        },
        'experiences': [
            {
                'company': 'Tech Solutions Inc.',
                'position': 'Senior Software Engineer',
                'location': 'New York, NY',
                'start_date': '2018-01-01',
                'end_date': None,
                'current': True,
                'description': 'Led development of cloud-based applications...'
            },
            {
                'company': 'Digital Innovations',
                'position': 'Software Developer',
                'location': 'Boston, MA',
                'start_date': '2015-03-01',
                'end_date': '2017-12-31',
                'current': False,
                'description': 'Developed and maintained web applications...'
            }
        ],
        'educations': [
            {
                'institution': 'MIT',
                'degree': 'Master of Science',
                'field_of_study': 'Computer Science',
                'location': 'Cambridge, MA',
                'start_date': '2013-09-01',
                'end_date': '2015-05-31',
                'current': False,
                'description': ''
            }
        ],
        'skills': [
            {'name': 'JavaScript', 'level': 5},
            {'name': 'Python', 'level': 4},
            {'name': 'React', 'level': 4},
            {'name': 'Node.js', 'level': 3},
            {'name': 'SQL', 'level': 4}
        ],
        'languages': [
            {'name': 'English', 'proficiency': 'native'},
            {'name': 'Spanish', 'proficiency': 'professional'}
        ],
        'projects': [
            {
                'title': 'E-commerce Platform',
                'description': 'Developed a full-stack e-commerce platform...',
                'url': 'https://github.com/johndoe/ecommerce',
                'start_date': '2019-06-01',
                'end_date': '2019-12-31'
            }
        ],
        'certifications': [
            {
                'name': 'AWS Certified Solutions Architect',
                'issuing_organization': 'Amazon Web Services',
                'issue_date': '2020-01-15',
                'expiration_date': '2023-01-15',
                'credential_id': 'AWS-123456',
                'credential_url': 'https://aws.amazon.com/certification/verify'
            }
        ]
    }

    # Get related templates (same category)
    related_templates = ResumeTemplate.objects.filter(
        category=template.category,
        is_active=True
    ).exclude(id=template.id).order_by('?')[:3]

    # Get color schemes for customization preview
    color_schemes = [
        {'name': 'Default', 'primary': '#3b82f6', 'secondary': '#6c757d', 'accent': '#10b981'},
        {'name': 'Professional', 'primary': '#1e40af', 'secondary': '#475569', 'accent': '#0369a1'},
        {'name': 'Creative', 'primary': '#7c3aed', 'secondary': '#6b7280', 'accent': '#ec4899'},
        {'name': 'Modern', 'primary': '#0f766e', 'secondary': '#4b5563', 'accent': '#f59e0b'},
        {'name': 'Classic', 'primary': '#b91c1c', 'secondary': '#374151', 'accent': '#ca8a04'}
    ]

    return render(request, 'templates_app/template_detail.html', {
        'template': template,
        'sample_data': sample_data,
        'related_templates': related_templates,
        'color_schemes': color_schemes
    })

# views.py
def template_gallery(request):
    categories = TemplateCategory.objects.all()
    selected_category = request.GET.get('category')
    search_query = request.GET.get('search', '')

    templates = ResumeTemplate.objects.filter(is_active=True)

    # Apply filters
    if selected_category:
        category = get_object_or_404(TemplateCategory, slug=selected_category)
        templates = templates.filter(category=category)

    if search_query:
        templates = templates.filter(
            Q(name__icontains=search_query) |
            Q(description__icontains=search_query) |
            Q(category__name__icontains=search_query)
        )

    # Get featured templates (one from each category)
    featured_templates = []
    for category in categories:
        featured = ResumeTemplate.objects.filter(
            category=category,
            is_active=True
        ).order_by('?').first()
        if featured:
            featured_templates.append(featured)

    return render(request, 'templates_app/template_gallery.html', {
        'categories': categories,
        'selected_category': selected_category,
        'search_query': search_query,
        'templates': templates,
        'featured_templates': featured_templates
    })

def home(request):
    """Landing page view"""
    # Count resumes created today
    from django.utils import timezone
    from django.db.models import Count
    import random

    today = timezone.now().date()
    # For demo purposes, generate a random number or use actual count
    resumes_today = Resume.objects.filter(created_at__date=today).count()
    if resumes_today < 100:  # If it's a new site or few users
        resumes_today = random.randint(1000, 50000)  # Show an impressive number

    featured_templates = ResumeTemplate.objects.filter(is_featured=True)[:3]

    return render(request, 'home/index.html', {
        'resumes_today': resumes_today,
        'featured_templates': featured_templates,
    })


@login_required
def my_resumes(request):
    """View for displaying user's resumes"""
    resumes = Resume.objects.filter(user=request.user).order_by('-updated_at')

    return render(request, 'dashboard/my_resumes.html', {
        'resumes': resumes
    })


class ResumeTemplateListView(ListView):
    """View for browsing resume templates"""
    model = ResumeTemplate
    template_name = 'dashboard/resume_list.html'
    context_object_name = 'templates'

    def get_queryset(self):
        queryset = super().get_queryset()
        category = self.request.GET.get('category')
        if category:
            queryset = queryset.filter(category=category)
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['categories'] = dict(ResumeTemplate.CATEGORY_CHOICES)
        context['selected_category'] = self.request.GET.get('category', '')
        return context


def resume_download(request, resume_id):
    resume = get_object_or_404(Resume, id=resume_id)

    # Set the context for the template
    context = {
        'resume': resume,
        'active_tab': 'download',
    }

    # If the user is just viewing the download page
    if request.method == 'GET':
        return render(request, 'templates_app/resume_download.html', context)

    # If the user has submitted the form to actually download the PDF
    elif request.method == 'POST':
        # Get the selected format
        selected_format = request.POST.get('format', 'pdf')

        if selected_format == 'pdf':
            # Create a PDF from the resume
            template = get_template('templates_app/resume_pdf_template.html')
            html = template.render({'resume': resume})

            # Create a file-like buffer to receive PDF data
            buffer = io.BytesIO()

            # Convert HTML to PDF
            pisa_status = pisa.CreatePDF(html, dest=buffer)

            # If error creating PDF
            if pisa_status.err:
                return HttpResponse('We had some errors with creating the PDF <pre>' + html + '</pre>')

            # FileResponse sets the Content-Disposition header so that browsers
            # present the option to save the file.
            buffer.seek(0)

            # Create the HTTP response with PDF
            response = HttpResponse(buffer, content_type='application/pdf')
            filename = f"{resume.title.replace(' ', '_')}_resume.pdf"
            response['Content-Disposition'] = f'attachment; filename="{filename}"'

            return response

        elif selected_format == 'docx':
            # This is a placeholder for DOCX generation
            # You would implement actual DOCX generation here
            response = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            filename = f"{resume.title.replace(' ', '_')}_resume.docx"
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            response.write(b'DOCX content would go here')  # Placeholder
            return response

        elif selected_format == 'txt':
            # Simple text version of the resume
            # Since we don't know the exact structure, we'll create a basic text file
            # with the information we do know
            content = f"""
RESUME: {resume.title}
Created: {resume.created_at.strftime('%B %d, %Y')}
Last Updated: {resume.updated_at.strftime('%B %d, %Y')}

"""

            # Try to get personal information if it exists
            try:
                # Check if there's a related personal_info model
                if hasattr(resume, 'personal_info'):
                    personal_info = resume.personal_info
                    content += "PERSONAL INFORMATION\n"
                    for field in personal_info._meta.fields:
                        if field.name not in ['id', 'resume', 'created_at', 'updated_at']:
                            value = getattr(personal_info, field.name, '')
                            if value:
                                content += f"{field.verbose_name}: {value}\n"
                    content += "\n"
            except Exception as e:
                # If there's an error, just continue without this section
                pass

            # Try to get sections if they exist
            try:
                # Check if there's a related sections queryset
                if hasattr(resume, 'sections'):
                    sections = resume.sections.all()
                    for section in sections:
                        section_title = getattr(section, 'title', 'Section')
                        content += f"{section_title.upper()}\n"

                        # Try to get section content
                        section_content = getattr(section, 'content', '')
                        if section_content:
                            content += f"{section_content}\n"

                        # Try to get items if this is a list-type section
                        try:
                            if hasattr(section, 'items'):
                                items = section.items.all()
                                for item in items:
                                    # Try common field names for items
                                    for field_name in ['title', 'subtitle', 'date_range', 'description']:
                                        value = getattr(item, field_name, '')
                                        if value:
                                            content += f"{value}\n"
                                    content += "\n"
                        except Exception:
                            pass

                        content += "\n"
            except Exception as e:
                # If there's an error, just continue without this section
                pass

            # Add a note about the resume design
            content += f"""
DESIGN INFORMATION
Template: {resume.template.name if resume.template else 'Custom'}
Primary Color: {resume.primary_color}
Font Family: {resume.get_font_family_display()}
Font Size: {resume.get_font_size_display()}
"""

            response = HttpResponse(content, content_type='text/plain')
            filename = f"{resume.title.replace(' ', '_')}_resume.txt"
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response

        else:
            # Default case - return to the download page with an error message
            context['error'] = f"Unsupported format: {selected_format}"
            return render(request, 'templates_app/resume_download.html', context)


def resume_design(request, resume_id):
    resume = get_object_or_404(Resume, id=resume_id)
    # Add any other context data you need
    context = {
        'resume': resume,
        'active_tab': 'design',
    }
    return render(request, 'templates_app/resume_design.html', context)


@login_required
def get_section(request, section_id):
    """Get section data for editing"""
    section = get_object_or_404(ResumeSection, id=section_id)

    # Check if user owns this section
    if section.resume.user != request.user:
        return JsonResponse({'error': 'Unauthorized'}, status=403)

    # Generate form HTML based on section type
    form_html = generate_section_form(section)

    return JsonResponse({
        'section': {
            'id': section.id,
            'name': section.name,
            'content': section.content
        },
        'form_html': form_html
    })


@login_required
@require_POST
def update_section(request, section_id):
    """Update section content"""
    section = get_object_or_404(ResumeSection, id=section_id)

    # Check if user owns this section
    if section.resume.user != request.user:
        return JsonResponse({'error': 'Unauthorized'}, status=403)

    # Process form data based on section type
    if section.name == 'Contact Information':
        section.content = {
            'name': request.POST.get('name', ''),
            'email': request.POST.get('email', ''),
            'phone': request.POST.get('phone', ''),
            'location': request.POST.get('location', '')
        }
    elif section.name == 'Professional Summary':
        section.content = {
            'summary': request.POST.get('summary', '')
        }
    elif section.name == 'Work Experience':
        # Process multiple job entries
        jobs = []
        job_count = int(request.POST.get('job_count', 0))

        for i in range(job_count):
            jobs.append({
                'title': request.POST.get(f'job_title_{i}', ''),
                'company': request.POST.get(f'job_company_{i}', ''),
                'dates': request.POST.get(f'job_dates_{i}', ''),
                'description': request.POST.get(f'job_description_{i}', '')
            })

        section.content = {'jobs': jobs}
    elif section.name == 'Education':
        # Process multiple education entries
        education = []
        edu_count = int(request.POST.get('edu_count', 0))

        for i in range(edu_count):
            education.append({
                'degree': request.POST.get(f'edu_degree_{i}', ''),
                'institution': request.POST.get(f'edu_institution_{i}', ''),
                'dates': request.POST.get(f'edu_dates_{i}', ''),
                'description': request.POST.get(f'edu_description_{i}', '')
            })

        section.content = {'education': education}
    elif section.name == 'Skills':
        # Process skills as a list
        skills_text = request.POST.get('skills', '')
        skills = [skill.strip() for skill in skills_text.split(',') if skill.strip()]
        section.content = {'skills': skills}
    else:
        # Generic section
        section.content = {
            'text': request.POST.get('content', '')
        }

    section.save()

    return JsonResponse({'success': True})


@login_required
@require_POST
def update_section_order(request, resume_id):
    """Update the order of resume sections"""
    resume = get_object_or_404(Resume, id=resume_id, user=request.user)

    section_order = json.loads(request.POST.get('section_order', '[]'))

    # Update order for each section
    for index, section_id in enumerate(section_order):
        section = get_object_or_404(ResumeSection, id=section_id, resume=resume)
        section.order = index + 1
        section.save()

    return JsonResponse({'success': True})


@login_required
@require_POST
def update_color_scheme(request, resume_id):
    """Update resume color scheme"""
    resume = get_object_or_404(Resume, id=resume_id, user=request.user)

    color_id = request.POST.get('color_id')
    color_scheme = get_object_or_404(TemplateColor, id=color_id, template=resume.template)

    resume.color_scheme = color_scheme
    resume.save()

    return JsonResponse({
        'success': True,
        'primary_color': color_scheme.primary_color,
        'secondary_color': color_scheme.secondary_color
    })


@login_required
@require_POST
def save_resume(request, resume_id):
    """Save resume"""
    resume = get_object_or_404(Resume, id=resume_id, user=request.user)

    # Update last_updated timestamp
    resume.save()

    return JsonResponse({'success': True})


@login_required
def preview_resume(request, resume_id):
    """Preview resume"""
    resume = get_object_or_404(Resume, id=resume_id, user=request.user)

    # Get all sections for this resume
    sections = resume.sections.all().order_by('order')

    return render(request, 'templates_app/resume_preview.html', {
        'resume': resume,
        'template': resume.template,
        'sections': sections,
    })


def generate_section_form(section):
    """Generate HTML form for editing a section based on its type"""
    if section.name == 'Contact Information':
        return f"""
        <form id="section-form">
            <input type="hidden" name="csrfmiddlewaretoken" value="">
            <div class="form-group">
                <label for="name">Full Name</label>
                <input type="text" id="name" name="name" class="form-control" 
                       value="{section.content.get('name', '')}">
            </div>
            <div class="form-group">
                <label for="email">Email</label>
                <input type="email" id="email" name="email" class="form-control"
                       value="{section.content.get('email', '')}">
            </div>
            <div class="form-group">
                <label for="phone">Phone</label>
                <input type="text" id="phone" name="phone" class="form-control"
                       value="{section.content.get('phone', '')}">
            </div>
            <div class="form-group">
                <label for="location">Location</label>
                <input type="text" id="location" name="location" class="form-control"
                       value="{section.content.get('location', '')}">
            </div>
        </form>
        """
    elif section.name == 'Professional Summary':
        return f"""
        <form id="section-form">
            <input type="hidden" name="csrfmiddlewaretoken" value="">
            <div class="form-group">
                <label for="summary">Professional Summary</label>
                <textarea id="summary" name="summary" class="form-control" rows="5">{section.content.get('summary', '')}</textarea>
            </div>
        </form>
        """
    elif section.name == 'Work Experience':
        jobs = section.content.get('jobs', [])
        if not jobs:
            jobs = [{'title': '', 'company': '', 'dates': '', 'description': ''}]

        jobs_html = ""
        for i, job in enumerate(jobs):
            jobs_html += f"""
            <div class="job-entry-form" data-index="{i}">
                <h4>Job #{i + 1}</h4>
                <div class="form-group">
                    <label for="job_title_{i}">Job Title</label>
                    <input type="text" id="job_title_{i}" name="job_title_{i}" class="form-control" 
                           value="{job.get('title', '')}">
                </div>
                <div class="form-group">
                    <label for="job_company_{i}">Company</label>
                    <input type="text" id="job_company_{i}" name="job_company_{i}" class="form-control"
                           value="{job.get('company', '')}">
                </div>
                <div class="form-group">
                    <label for="job_dates_{i}">Dates</label>
                    <input type="text" id="job_dates_{i}" name="job_dates_{i}" class="form-control"
                           value="{job.get('dates', '')}">
                </div>
                <div class="form-group">
                    <label for="job_description_{i}">Description</label>
                    <textarea id="job_description_{i}" name="job_description_{i}" class="form-control" rows="3">{job.get('description', '')}</textarea>
                </div>
                <button type="button" class="btn btn-outline btn-sm remove-job">Remove</button>
            </div>
            """

        return f"""
        <form id="section-form">
            <input type="hidden" name="csrfmiddlewaretoken" value="">
            <input type="hidden" name="job_count" value="{len(jobs)}" id="job_count">
            <div id="jobs-container">
                {jobs_html}
            </div>
            <button type="button" id="add-job" class="btn btn-outline btn-sm">Add Another Job</button>
        </form>
        <script>
            $(document).ready(function() {{
                // Add job
                $("#add-job").click(function() {{
                    const index = $(".job-entry-form").length;
                    const newJob = `
                        <div class="job-entry-form" data-index="${{index}}">
                            <h4>Job #${{index+1}}</h4>
                            <div class="form-group">
                                <label for="job_title_${{index}}">Job Title</label>
                                <input type="text" id="job_title_${{index}}" name="job_title_${{index}}" class="form-control">
                            </div>
                            <div class="form-group">
                                <label for="job_company_${{index}}">Company</label>
                                <input type="text" id="job_company_${{index}}" name="job_company_${{index}}" class="form-control">
                            </div>
                            <div class="form-group">
                                <label for="job_dates_${{index}}">Dates</label>
                                <input type="text" id="job_dates_${{index}}" name="job_dates_${{index}}" class="form-control">
                            </div>
                            <div class="form-group">
                                <label for="job_description_${{index}}">Description</label>
                                <textarea id="job_description_${{index}}" name="job_description_${{index}}" class="form-control" rows="3"></textarea>
                            </div>
                            <button type="button" class="btn btn-outline btn-sm remove-job">Remove</button>
                        </div>
                    `;
                    $("#jobs-container").append(newJob);
                    $("#job_count").val(index + 1);
                }});

                // Remove job
                $(document).on("click", ".remove-job", function() {{
                    $(this).closest(".job-entry-form").remove();
                    // Update job count
                    $("#job_count").val($(".job-entry-form").length);
                    // Renumber jobs
                    $(".job-entry-form").each(function(index) {{
                        $(this).find("h4").text(`Job #${{index+1}}`);
                    }});
                }});
            }});
        </script>
        """
    # Add similar form generators for other section types

    # Default generic form
    return f"""
    <form id="section-form">
        <input type="hidden" name="csrfmiddlewaretoken" value="">
        <div class="form-group">
            <label for="content">Content</label>
            <textarea id="content" name="content" class="form-control" rows="5">{section.content.get('text', '')}</textarea>
        </div>
    </form>
    """


@login_required
def resume_editor(request):
    """Resume editor view"""
    # Get template from query parameter
    template_slug = request.GET.get('template')
    if not template_slug:
        messages.error(request, "No template selected.")
        return redirect('template_list')

    # Import the correct ResumeTemplate model
    from dashboard.models import ResumeTemplate

    # Get the template object using the slug
    try:
        template = ResumeTemplate.objects.get(slug=template_slug)
    except ResumeTemplate.DoesNotExist:
        messages.error(request, f"Template '{template_slug}' not found.")
        return redirect('templates_app:template_list')

    # Get or create a resume for this user and template
    try:
        # Make sure to query using the template instance, not a string
        resume = Resume.objects.get(
            user=request.user,
            template=template  # This should be a ResumeTemplate instance
        )
    except Resume.DoesNotExist:
        # Create a new resume with the template
        resume = Resume.objects.create(
            user=request.user,
            template=template,
            title=f"My {template.name} Resume",
            full_name=request.user.get_full_name() or request.user.username,
            email=request.user.email,
            primary_color=template.default_primary_color if hasattr(template, 'default_primary_color') else '#2D9CDB',
            secondary_color=template.default_secondary_color if hasattr(template,
                                                                        'default_secondary_color') else '#3498DB'
        )

    # Get sections data (assuming you have related models for sections)
    # Adjust this based on your actual models
    education_entries = resume.education_entries.all() if hasattr(resume, 'education_entries') else []
    experience_entries = resume.experience_entries.all() if hasattr(resume, 'experience_entries') else []
    skill_entries = resume.skill_entries.all() if hasattr(resume, 'skill_entries') else []

    return render(request, 'template_list/resume_editor.html', {
        'resume': resume,
        'template': template,
        'education_entries': education_entries,
        'experience_entries': experience_entries,
        'skill_entries': skill_entries,
    })


class TemplateListView(ListView):
    """Display all resume templates with filtering options"""
    model = ResumeTemplate
    template_name = 'templates_app/template_list.html'
    context_object_name = 'templates'
    paginate_by = 12

    def get_queryset(self):
        queryset = ResumeTemplate.objects.all()

        # Filter by category if provided
        category_slug = self.request.GET.get('category')
        if category_slug:
            queryset = queryset.filter(categories__slug=category_slug)

        # Filter by ATS-friendly
        ats_friendly = self.request.GET.get('ats_friendly')
        if ats_friendly == 'true':
            queryset = queryset.filter(is_ats_friendly=True)

        # Filter by photo option
        has_photo = self.request.GET.get('has_photo')
        if has_photo == 'true':
            queryset = queryset.filter(has_photo=True)
        elif has_photo == 'false':
            queryset = queryset.filter(has_photo=False)

        # Filter by premium status
        premium = self.request.GET.get('premium')
        if premium == 'true':
            queryset = queryset.filter(is_premium=True)
        elif premium == 'false':
            queryset = queryset.filter(is_premium=False)

        # Search by name
        search_query = self.request.GET.get('q')
        if search_query:
            queryset = queryset.filter(
                Q(name__icontains=search_query) |
                Q(description__icontains=search_query)
            )

        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['categories'] = TemplateCategory.objects.all()
        context['filter_form'] = TemplateFilterForm(self.request.GET)

        # Get active category for highlighting
        category_slug = self.request.GET.get('category')
        if category_slug:
            context['active_category'] = get_object_or_404(TemplateCategory, slug=category_slug)

        return context


class TemplateDetailView(DetailView):
    """Display details of a specific template"""
    model = ResumeTemplate
    template_name = 'templates_app/template_detail.html'
    context_object_name = 'template'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        template = self.get_object()

        # Get color schemes
        context['colors'] = template.colors.all()

        # Get default color scheme
        context['default_color'] = template.colors.filter(is_default=True).first()

        # Get template sections
        context['sections'] = template.sections.all()

        # Check if user has used this template before
        if self.request.user.is_authenticated:
            context['user_has_used'] = UserTemplateSelection.objects.filter(
                user=self.request.user,
                template=template
            ).exists()

        return context


@login_required
def select_template(request, slug):
    """Select a template to create a resume"""
    template = get_object_or_404(ResumeTemplate, slug=slug)

    # Check if template is premium and user has access
    if template.is_premium and not request.user.profile.has_premium_access:
        messages.error(request, "This is a premium template. Please upgrade to access it.")
        return redirect('templates_app:template_list')

    # Get color scheme if provided
    color_scheme = None
    color_id = request.POST.get('color_scheme')
    if color_id:
        color_scheme = get_object_or_404(TemplateColor, id=color_id, template=template)
    else:
        # Use default color scheme
        color_scheme = template.colors.filter(is_default=True).first()

    # Create or update user template selection
    selection, created = UserTemplateSelection.objects.update_or_create(
        user=request.user,
        template=template,
        defaults={'color_scheme': color_scheme}
    )

    # Increment template popularity
    template.increment_popularity()

    # Redirect to resume editor with the selected template
    return redirect(reverse('templates_app:resume_editor') + f'?template={template.slug}')


def template_preview(request, slug):
    """Preview a template with sample data"""
    template = get_object_or_404(ResumeTemplate, slug=slug)

    # Get color scheme if provided
    color_id = request.GET.get('color')
    if color_id:
        color = get_object_or_404(TemplateColor, id=color_id, template=template)
    else:
        # Use default color scheme
        color = template.colors.filter(is_default=True).first()

    context = {
        'template': template,
        'color': color,
        'is_preview': True,
    }

    return render(request, 'templates_app/template_preview.html', context)


def template_filter_ajax(request):
    """AJAX endpoint for filtering templates"""
    category_slug = request.GET.get('category')
    ats_friendly = request.GET.get('ats_friendly')
    has_photo = request.GET.get('has_photo')
    premium = request.GET.get('premium')
    search_query = request.GET.get('q')

    templates = ResumeTemplate.objects.all()

    # Apply filters
    if category_slug:
        templates = templates.filter(categories__slug=category_slug)

    if ats_friendly == 'true':
        templates = templates.filter(is_ats_friendly=True)

    if has_photo == 'true':
        templates = templates.filter(has_photo=True)
    elif has_photo == 'false':
        templates = templates.filter(has_photo=False)

    if premium == 'true':
        templates = templates.filter(is_premium=True)
    elif premium == 'false':
        templates = templates.filter(is_premium=False)

    if search_query:
        templates = templates.filter(
            Q(name__icontains=search_query) |
            Q(description__icontains=search_query)
        )

    # Render template list partial
    html = render(request, 'templates_app/partials/template_grid.html', {
        'templates': templates
    }).content.decode('utf-8')

    return JsonResponse({
        'html': html,
        'count': templates.count()
    })


def subscription_plans(request):
    # Sample subscription plans data
    plans = [
        {
            'name': 'Basic',
            'price': '$9.99',
            'period': 'month',
            'features': [
                'Access to basic templates',
                'Email support',
                'Up to 3 projects',
                '1GB storage'
            ],
            'is_popular': False,
            'button_text': 'Get Started'
        },
        {
            'name': 'Pro',
            'price': '$19.99',
            'period': 'month',
            'features': [
                'Access to all templates',
                'Priority email support',
                'Up to 10 projects',
                '10GB storage',
                'Custom branding'
            ],
            'is_popular': True,
            'button_text': 'Get Pro'
        },
        {
            'name': 'Enterprise',
            'price': '$49.99',
            'period': 'month',
            'features': [
                'Access to all templates',
                'Phone and email support',
                'Unlimited projects',
                '100GB storage',
                'Custom branding',
                'API access',
                'Dedicated account manager'
            ],
            'is_popular': False,
            'button_text': 'Contact Sales'
        }
    ]

    return render(request, 'templates_app/contemporary/subscription_plans.html', {
        'plans': plans,
        'title': 'Subscription Plans'
    })


class ProfessionalTemplatesView(ListView):
    """View for browsing professional resume templates"""
    model = ResumeTemplate
    template_name = 'templates_app/professional_templates/professional_template.html'
    context_object_name = 'templates'

    def get_queryset(self):
        queryset = super().get_queryset().filter(category='professional')

        # Filter by style if provided
        style = self.request.GET.get('style')
        if style:
            queryset = queryset.filter(style=style)

        # Filter by color if provided
        color = self.request.GET.get('color')
        if color:
            queryset = queryset.filter(color=color)

        # Sort by popularity, newest, or name
        sort = self.request.GET.get('sort', 'popular')
        if sort == 'newest':
            queryset = queryset.order_by('-created_at')
        elif sort == 'name':
            queryset = queryset.order_by('name')
        else:  # Default to popular
            queryset = queryset.order_by('-popularity')

        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        # Add filter options to context
        context['styles'] = ResumeTemplate.objects.filter(category='professional').values_list('style',
                                                                                               flat=True).distinct()
        context['colors'] = ResumeTemplate.objects.filter(category='professional').values_list('color',
                                                                                               flat=True).distinct()

        # Add current filter selections
        context['selected_style'] = self.request.GET.get('style', '')
        context['selected_color'] = self.request.GET.get('color', '')
        context['selected_sort'] = self.request.GET.get('sort', 'popular')

        return context


@login_required
def create_resume(request, template_slug=None):
    """Create a new resume, optionally based on a template"""
    if template_slug:
        template = get_object_or_404(ResumeTemplate, slug=template_slug)
    else:
        # If no template specified, use the first simple template
        template = ResumeTemplate.objects.filter(category='simple').first()
        if not template:
            # Fallback to any template if no simple ones exist
            template = ResumeTemplate.objects.first()

    # Create a new resume with the selected template
    resume = Resume.objects.create(
        user=request.user,
        template=template,
        title=f"Untitled Resume - {template.name}",
        content={
            "personal_info": {
                "first_name": "",
                "last_name": "",
                "email": request.user.email,
                "phone": "",
                "address": "",
                "city": "",
                "state": "",
                "zip_code": "",
                "country": "",
                "linkedin": "",
                "website": ""
            }
        }
    )

    # Create default sections
    default_sections = [
        {"type": "summary", "title": "Professional Summary", "order": 1},
        {"type": "experience", "title": "Work Experience", "order": 2},
        {"type": "education", "title": "Education", "order": 3},
        {"type": "skills", "title": "Skills", "order": 4},
    ]

    for section in default_sections:
        ResumeSection.objects.create(
            resume=resume,
            section_type=section["type"],
            title=section["title"],
            order=section["order"],
            content={}
        )

    return redirect('dashboard:resume_edit', uuid=resume.uuid)


@login_required
def resume_edit(request, uuid):
    """Edit a resume"""
    resume = get_object_or_404(Resume, uuid=uuid, user=request.user)
    sections = resume.sections.all().order_by('order')

    if request.method == 'POST':
        form = ResumeForm(request.POST, instance=resume)
        if form.is_valid():
            form.save()
            messages.success(request, "Resume updated successfully!")
            return redirect('dashboard:resume_edit', uuid=resume.uuid)
    else:
        form = ResumeForm(instance=resume)

    return render(request, 'dashboard/resume_edit.html', {
        'resume': resume,
        'sections': sections,
        'form': form,
    })


@login_required
def resume_preview(request, uuid):
    """Preview a resume"""
    resume = get_object_or_404(Resume, uuid=uuid, user=request.user)
    sections = resume.sections.all().order_by('order')

    return render(request, 'dashboard/resume_preview.html', {
        'resume': resume,
        'sections': sections,
    })


@login_required
def upload_resume(request):
    """Upload an existing resume for parsing"""
    if request.method == 'POST':
        form = ResumeUploadForm(request.POST, request.FILES)
        if form.is_valid():
            # Here you would implement resume parsing logic
            # For now, we'll just create a blank resume
            template = ResumeTemplate.objects.filter(category='simple').first()

            resume = Resume.objects.create(
                user=request.user,
                template=template,
                title="Uploaded Resume",
                content={"personal_info": {"email": request.user.email}}
            )

            # Create default sections
            default_sections = [
                {"type": "summary", "title": "Professional Summary", "order": 1},
                {"type": "experience", "title": "Work Experience", "order": 2},
                {"type": "education", "title": "Education", "order": 3},
                {"type": "skills", "title": "Skills", "order": 4},
            ]

            for section in default_sections:
                ResumeSection.objects.create(
                    resume=resume,
                    section_type=section["type"],
                    title=section["title"],
                    order=section["order"],
                    content={}
                )

            messages.success(request, "Resume uploaded successfully! Please review and edit the extracted information.")
            return redirect('dashboard:resume_edit', uuid=resume.uuid)
    else:
        form = ResumeUploadForm()

    return render(request, 'dashboard/resume_upload.html', {'form': form})


@login_required
def add_section(request, resume_uuid):
    """AJAX endpoint to add a new section to a resume"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            resume = get_object_or_404(Resume, uuid=resume_uuid, user=request.user)

            # Get the highest order value and add 1
            highest_order = ResumeSection.objects.filter(resume=resume).order_by('-order').first()
            new_order = 1 if not highest_order else highest_order.order + 1

            section = ResumeSection.objects.create(
                resume=resume,
                section_type=data.get('section_type', 'custom'),
                title=data.get('title', 'New Section'),
                order=new_order,
                content={}
            )

            return JsonResponse({
                'success': True,
                'section_id': section.id,
                'section_type': section.section_type,
                'title': section.title
            })
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})

    return JsonResponse({'success': False, 'error': 'Invalid request method'})


@login_required
def save_resume(request, uuid):
    """AJAX endpoint to save resume data"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            resume = get_object_or_404(Resume, uuid=uuid, user=request.user)

            # Update personal info
            if 'personal_info' in data:
                resume.content['personal_info'] = data['personal_info']

            # Update sections
            if 'sections' in data:
                for section_data in data['sections']:
                    section_id = section_data.get('id')
                    section = get_object_or_404(ResumeSection, id=section_id, resume=resume)

                    # Update section content
                    section.content = section_data.get('content', {})
                    section.save()

            resume.save()

            return JsonResponse({'success': True})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})

    return JsonResponse({'success': False, 'error': 'Invalid request method'})


@login_required
def delete_resume(request, resume_id):
    """Delete a resume"""
    resume = get_object_or_404(Resume, id=resume_id, user=request.user)

    if request.method == 'POST':
        resume_title = resume.title
        resume.delete()
        messages.success(request, f"Resume '{resume_title}' has been deleted.")
        return redirect('dashboard:my_resumes')

    return redirect('dashboard:my_resumes')


@login_required
def download_resume(request, uuid, format):
    """Download resume in various formats"""
    resume = get_object_or_404(Resume, uuid=uuid, user=request.user)
    sections = resume.sections.all().order_by('order')

    if format == 'pdf':
        # For PDF generation, you would typically use a library like WeasyPrint or xhtml2pdf
        # This is a simplified example
        from django.template.loader import render_to_string
        from django.http import HttpResponse
        from weasyprint import HTML, CSS
        from django.conf import settings
        import tempfile

        # Render the resume to HTML
        html_string = render_to_string('dashboard/resume_pdf.html', {
            'resume': resume,
            'sections': sections,
            'base_url': request.build_absolute_uri('/'),
        })

        # Create HTTP response with PDF
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{resume.title}.pdf"'

        # Generate PDF from HTML
        HTML(string=html_string, base_url=request.build_absolute_uri('/')).write_pdf(
            response,
            stylesheets=[CSS(string=resume.template.css_styles)]
        )

        return response

    elif format == 'docx':
        # For DOCX generation, you would typically use a library like python-docx
        # This is a simplified example
        from django.http import HttpResponse
        import docx
        from docx.shared import Pt

        # Create a new document
        doc = docx.Document()

        # Add title
        doc.add_heading(
            f"{resume.content.get('personal_info', {}).get('first_name', '')} {resume.content.get('personal_info', {}).get('last_name', '')}",
            0)

        # Add contact info
        contact_info = resume.content.get('personal_info', {})
        p = doc.add_paragraph()
        p.add_run(f"Email: {contact_info.get('email', '')}\n")
        p.add_run(f"Phone: {contact_info.get('phone', '')}\n")
        p.add_run(
            f"Address: {contact_info.get('address', '')}, {contact_info.get('city', '')}, {contact_info.get('state', '')} {contact_info.get('zip_code', '')}")

        # Add sections
        for section in sections:
            doc.add_heading(section.title, 1)

            if section.section_type == 'summary':
                doc.add_paragraph(section.content.get('summary', ''))

            elif section.section_type == 'experience':
                for item in section.content.get('items', []):
                    p = doc.add_paragraph()
                    p.add_run(f"{item.get('job_title', '')} at {item.get('employer', '')}").bold = True
                    p.add_run(
                        f"\n{item.get('start_date', '')} - {item.get('end_date', 'Present') if not item.get('current_job', False) else 'Present'}")
                    doc.add_paragraph(item.get('description', ''))

            elif section.section_type == 'education':
                for item in section.content.get('items', []):
                    p = doc.add_paragraph()
                    p.add_run(f"{item.get('degree', '')} - {item.get('institution', '')}").bold = True
                    p.add_run(
                        f"\n{item.get('start_date', '')} - {item.get('end_date', 'Present') if not item.get('current_education', False) else 'Present'}")
                    doc.add_paragraph(item.get('description', ''))

            elif section.section_type == 'skills':
                skills = section.content.get('skills', [])
                if skills:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(', '.join(skills))

            else:
                doc.add_paragraph(section.content.get('text', ''))

        # Create response
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{resume.title}.docx"'

        # Save document to response
        doc.save(response)

        return response

    elif format == 'txt':
        # For TXT generation
        from django.http import HttpResponse

        # Create response
        response = HttpResponse(content_type='text/plain')
        response['Content-Disposition'] = f'attachment; filename="{resume.title}.txt"'

        # Write personal info
        personal_info = resume.content.get('personal_info', {})
        response.write(f"{personal_info.get('first_name', '')} {personal_info.get('last_name', '')}\n")
        response.write(f"Email: {personal_info.get('email', '')}\n")
        response.write(f"Phone: {personal_info.get('phone', '')}\n")
        response.write(
            f"Address: {personal_info.get('address', '')}, {personal_info.get('city', '')}, {personal_info.get('state', '')} {personal_info.get('zip_code', '')}\n\n")

        # Write sections
        for section in sections:
            response.write(f"{section.title.upper()}\n")
            response.write("=" * len(section.title) + "\n\n")

            if section.section_type == 'summary':
                response.write(f"{section.content.get('summary', '')}\n\n")

            elif section.section_type == 'experience':
                for item in section.content.get('items', []):
                    response.write(f"{item.get('job_title', '')} at {item.get('employer', '')}\n")
                    response.write(
                        f"{item.get('start_date', '')} - {item.get('end_date', 'Present') if not item.get('current_job', False) else 'Present'}\n")
                    response.write(f"{item.get('description', '')}\n\n")

            elif section.section_type == 'education':
                for item in section.content.get('items', []):
                    response.write(f"{item.get('degree', '')} - {item.get('institution', '')}\n")
                    response.write(
                        f"{item.get('start_date', '')} - {item.get('end_date', 'Present') if not item.get('current_education', False) else 'Present'}\n")
                    response.write(f"{item.get('description', '')}\n\n")

            elif section.section_type == 'skills':
                skills = section.content.get('skills', [])
                if skills:
                    response.write(', '.join(skills) + "\n\n")

            else:
                response.write(f"{section.content.get('text', '')}\n\n")

        return response

    else:
        # Invalid format
        messages.error(request, f"Invalid download format: {format}")
        return redirect('dashboard:resume_preview', uuid=uuid)
