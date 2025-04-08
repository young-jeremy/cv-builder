from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.auth import login
from django.contrib import messages
from django.http import HttpResponse, JsonResponse
from django.template.loader import render_to_string
from django.urls import reverse
from django.forms import modelformset_factory
from django.db.models import Max
from django.views.decorators.http import require_POST
from django.core.paginator import Paginator
import json
import uuid
from weasyprint import HTML, CSS
from docx import Document
from io import BytesIO
from templates_app.models import ResumeTemplate

from .models import (
    TemplateCategory, ResumeTemplate, PersonalInfo,
    Experience, Education, Skill, Project, Language, Certification, TemplateColor
)
from .forms import (
    UserRegistrationForm, ResumeForm, PersonalInfoForm, ExperienceForm,
    EducationForm, SkillForm, ProjectForm, LanguageForm, CertificationForm,
    TemplateColorForm, ResumeExportForm
)


@login_required
@require_POST
def add_project(request, resume_id):
    resume = get_object_or_404(ResumeTemplate, uuid=resume_id, user=request.user)
    form = ProjectForm(request.POST)

    if form.is_valid():
        project = form.save(commit=False)
        project.resume = resume

        # Set order to be the highest + 1
        max_order = resume.projects.aggregate(Max('order'))['order__max'] or 0
        project.order = max_order + 1

        project.save()
        messages.success(request, 'Project added successfully!')
    else:
        messages.error(request, 'Error adding project.')

    return redirect('edit_resume', resume_id=resume.uuid)


@login_required
@require_POST
def update_project(request, resume_id, project_id):
    resume = get_object_or_404(ResumeTemplate, uuid=resume_id, user=request.user)
    project = get_object_or_404(Project, id=project_id, resume=resume)

    form = ProjectForm(request.POST, instance=project)

    if form.is_valid():
        form.save()
        messages.success(request, 'Project updated successfully!')
    else:
        messages.error(request, 'Error updating project.')

    return redirect('edit_resume', resume_id=resume.uuid)


@login_required
@require_POST
def delete_project(request, resume_id, project_id):
    resume = get_object_or_404(ResumeTemplate, uuid=resume_id, user=request.user)
    project = get_object_or_404(Project, id=project_id, resume=resume)

    project.delete()
    messages.success(request, 'Project deleted successfully!')

    return redirect('edit_resume', resume_id=resume.uuid)


@login_required
@require_POST
def add_language(request, resume_id):
    resume = get_object_or_404(ResumeTemplate, uuid=resume_id, user=request.user)
    form = LanguageForm(request.POST)

    if form.is_valid():
        language = form.save(commit=False)
        language.resume = resume

        # Set order to be the highest + 1
        max_order = resume.languages.aggregate(Max('order'))['order__max'] or 0
        language.order = max_order + 1

        language.save()
        messages.success(request, 'Language added successfully!')
    else:
        messages.error(request, 'Error adding language.')

    return redirect('edit_resume', resume_id=resume.uuid)


@login_required
@require_POST
def update_language(request, resume_id, language_id):
    resume = get_object_or_404(ResumeTemplate, uuid=resume_id, user=request.user)
    language = get_object_or_404(Language, id=language_id, resume=resume)

    form = LanguageForm(request.POST, instance=language)

    if form.is_valid():
        form.save()
        messages.success(request, 'Language updated successfully!')
    else:
        messages.error(request, 'Error updating language.')

    return redirect('edit_resume', resume_id=resume.uuid)


@login_required
@require_POST
def delete_language(request, resume_id, language_id):
    resume = get_object_or_404(ResumeTemplate, uuid=resume_id, user=request.user)
    language = get_object_or_404(Language, id=language_id, resume=resume)

    language.delete()
    messages.success(request, 'Language deleted successfully!')

    return redirect('edit_resume', resume_id=resume.uuid)


@login_required
@require_POST
def add_certification(request, resume_id):
    resume = get_object_or_404(ResumeTemplate, uuid=resume_id, user=request.user)
    form = CertificationForm(request.POST)

    if form.is_valid():
        certification = form.save(commit=False)
        certification.resume = resume

        # Set order to be the highest + 1
        max_order = resume.certifications.aggregate(Max('order'))['order__max'] or 0
        certification.order = max_order + 1

        certification.save()
        messages.success(request, 'Certification added successfully!')
    else:
        messages.error(request, 'Error adding certification.')

    return redirect('edit_resume', resume_id=resume.uuid)


@login_required
@require_POST
def update_certification(request, resume_id, certification_id):
    resume = get_object_or_404(ResumeTemplate, uuid=resume_id, user=request.user)
    certification = get_object_or_404(Certification, id=certification_id, resume=resume)

    form = CertificationForm(request.POST, instance=certification)

    if form.is_valid():
        form.save()
        messages.success(request, 'Certification updated successfully!')
    else:
        messages.error(request, 'Error updating certification.')

    return redirect('edit_resume', resume_id=resume.uuid)


@login_required
@require_POST
def delete_certification(request, resume_id, certification_id):
    resume = get_object_or_404(ResumeTemplate, uuid=resume_id, user=request.user)
    certification = get_object_or_404(Certification, id=certification_id, resume=resume)

    certification.delete()
    messages.success(request, 'Certification deleted successfully!')

    return redirect('edit_resume', resume_id=resume.uuid)


def home(request):
    template_categories = TemplateCategory.objects.all()
    resume_count = ResumeTemplate.objects.count()
    return render(request, 'resume_builder/home.html', {
        'template_categories': template_categories,
        'resume_count': resume_count
    })


def register(request):
    if request.method == 'POST':
        form = UserRegistrationForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            messages.success(request, 'Account created successfully!')
            return redirect('home')
    else:
        form = UserRegistrationForm()
    return render(request, 'resume_builder/register.html', {'form': form})


@login_required
def dashboard(request):
    resumes = ResumeTemplate.objects.filter(user=request.user).order_by('-updated_at')
    return render(request, 'resume_builder/dashboard.html', {'resumes': resumes})


def template_list(request):
    categories = TemplateCategory.objects.all()
    selected_category = request.GET.get('category')

    if selected_category:
        category = get_object_or_404(TemplateCategory, slug=selected_category)
        templates = ResumeTemplate.objects.filter(category=category, is_active=True)
    else:
        templates = ResumeTemplate.objects.filter(is_active=True)

    paginator = Paginator(templates, 12)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    return render(request, 'resume_builder/template_list.html', {
        'categories': categories,
        'selected_category': selected_category,
        'page_obj': page_obj
    })


@login_required
def create_resume(request, template_slug=None):
    if template_slug:
        template = get_object_or_404(ResumeTemplate, slug=template_slug, is_active=True)
    else:
        template = None

    if request.method == 'POST':
        form = ResumeForm(request.POST)
        if form.is_valid():
            resume = form.save(commit=False)
            resume.user = request.user
            if template:
                resume.template = template
            resume.save()

            # Create empty personal info
            PersonalInfo.objects.create(resume=resume)

            messages.success(request, 'Resume created successfully!')
            return redirect('edit_resume', resume_id=resume.uuid)
    else:
        initial_data = {}
        if template:
            initial_data['template'] = template.id
        form = ResumeForm(initial=initial_data)

    return render(request, 'resume/create_resume.html', {
        'form': form,
        'template': template
    })


@login_required
def edit_resume(request, resume_id):
    resume = get_object_or_404(ResumeTemplate, uuid=resume_id, user=request.user)

    # Get or create personal info
    try:
        personal_info = resume.personal_info
    except PersonalInfo.DoesNotExist:
        personal_info = PersonalInfo.objects.create(resume=resume)

    # Get related data
    experiences = resume.experiences.all().order_by('order')
    educations = resume.educations.all().order_by('order')
    skills = resume.skills.all().order_by('order')
    projects = resume.projects.all().order_by('order')
    languages = resume.languages.all().order_by('order')
    certifications = resume.certifications.all().order_by('order')
    colors = resume.colors.all()

    # Initialize forms
    personal_info_form = PersonalInfoForm(instance=personal_info)
    experience_form = ExperienceForm()
    education_form = EducationForm()
    skill_form = SkillForm()
    project_form = ProjectForm()
    language_form = LanguageForm()
    certification_form = CertificationForm()
    color_form = TemplateColorForm()
    export_form = ResumeExportForm()

    return render(request, 'resume_builder/edit_resume.html', {
        'resume': resume,
        'personal_info_form': personal_info_form,
        'experience_form': experience_form,
        'education_form': education_form,
        'skill_form': skill_form,
        'project_form': project_form,
        'language_form': language_form,
        'certification_form': certification_form,
        'color_form': color_form,
        'export_form': export_form,
        'experiences': experiences,
        'educations': educations,
        'skills': skills,
        'projects': projects,
        'languages': languages,
        'certifications': certifications,
        'colors': colors,
    })


@login_required
@require_POST
def update_personal_info(request, resume_id):
    resume = get_object_or_404(ResumeTemplate, uuid=resume_id, user=request.user)

    try:
        personal_info = resume.personal_info
    except PersonalInfo.DoesNotExist:
        personal_info = PersonalInfo(resume=resume)

    form = PersonalInfoForm(request.POST, request.FILES, instance=personal_info)

    if form.is_valid():
        form.save()
        messages.success(request, 'Personal information updated successfully!')
    else:
        messages.error(request, 'Error updating personal information.')

    return redirect('edit_resume', resume_id=resume.uuid)


@login_required
@require_POST
def add_experience(request, resume_id):
    resume = get_object_or_404(ResumeTemplate, uuid=resume_id, user=request.user)
    form = ExperienceForm(request.POST)

    if form.is_valid():
        experience = form.save(commit=False)
        experience.resume = resume

        # Set order to be the highest + 1
        max_order = resume.experiences.aggregate(Max('order'))['order__max'] or 0
        experience.order = max_order + 1

        experience.save()
        messages.success(request, 'Experience added successfully!')
    else:
        messages.error(request, 'Error adding experience.')

    return redirect('edit_resume', resume_id=resume.uuid)


@login_required
@require_POST
def update_experience(request, resume_id, experience_id):
    resume = get_object_or_404(ResumeTemplate, uuid=resume_id, user=request.user)
    experience = get_object_or_404(Experience, id=experience_id, resume=resume)

    form = ExperienceForm(request.POST, instance=experience)

    if form.is_valid():
        form.save()
        messages.success(request, 'Experience updated successfully!')
    else:
        messages.error(request, 'Error updating experience.')

    return redirect('edit_resume', resume_id=resume.uuid)


@login_required
@require_POST
def delete_experience(request, resume_id, experience_id):
    resume = get_object_or_404(ResumeTemplate, uuid=resume_id, user=request.user)
    experience = get_object_or_404(Experience, id=experience_id, resume=resume)

    experience.delete()
    messages.success(request, 'Experience deleted successfully!')

    return redirect('edit_resume', resume_id=resume.uuid)


@login_required
@require_POST
def add_education(request, resume_id):
    resume = get_object_or_404(ResumeTemplate, uuid=resume_id, user=request.user)
    form = EducationForm(request.POST)

    if form.is_valid():
        education = form.save(commit=False)
        education.resume = resume

        # Set order to be the highest + 1
        max_order = resume.educations.aggregate(Max('order'))['order__max'] or 0
        education.order = max_order + 1

        education.save()
        messages.success(request, 'Education added successfully!')
    else:
        messages.error(request, 'Error adding education.')

    return redirect('templates_app:edit_resume', resume_id=resume.uuid)


@login_required
@require_POST
def update_education(request, resume_id, education_id):
    resume = get_object_or_404(ResumeTemplate, uuid=resume_id, user=request.user)
    education = get_object_or_404(Education, id=education_id, resume=resume)

    form = EducationForm(request.POST, instance=education)

    if form.is_valid():
        form.save()
        messages.success(request, 'Education updated successfully!')
    else:
        messages.error(request, 'Error updating education.')

    return redirect('edit_resume', resume_id=resume.uuid)


@login_required
@require_POST
def delete_education(request, resume_id, education_id):
    resume = get_object_or_404(ResumeTemplate, uuid=resume_id, user=request.user)
    education = get_object_or_404(Education, id=education_id, resume=resume)

    education.delete()
    messages.success(request, 'Education deleted successfully!')

    return redirect('edit_resume', resume_id=resume.uuid)


@login_required
@require_POST
def add_skill(request, resume_id):
    resume = get_object_or_404(ResumeTemplate, uuid=resume_id, user=request.user)
    form = SkillForm(request.POST)

    if form.is_valid():
        skill = form.save(commit=False)
        skill.resume = resume

        # Set order to be the highest + 1
        max_order = resume.skills.aggregate(Max('order'))['order__max'] or 0
        skill.order = max_order + 1

        skill.save()
        messages.success(request, 'Skill added successfully!')
    else:
        messages.error(request, 'Error adding skill.')

    return redirect('edit_resume', resume_id=resume.uuid)


@login_required
@require_POST
def update_skill(request, resume_id, skill_id):
    resume = get_object_or_404(ResumeTemplate, uuid=resume_id, user=request.user)
    skill = get_object_or_404(Skill, id=skill_id, resume=resume)

    form = SkillForm(request.POST, instance=skill)

    if form.is_valid():
        form.save()
        messages.success(request, 'Skill updated successfully!')
    else:
        messages.error(request, 'Error updating skill.')

    return redirect('templates_app:edit_resume', resume_id=resume.uuid)


@login_required
@require_POST
def delete_skill(request, resume_id, skill_id):
    resume = get_object_or_404(ResumeTemplate, uuid=resume_id, user=request.user)
    skill = get_object_or_404(Skill, id=skill_id, resume=resume)

    skill.delete()
    messages.success(request, 'Skill deleted successfully!')

    return redirect('templates_app:edit_resume', resume_id=resume.uuid)


# Similar CRUD views for Project, Language, and Certification
# (Omitted for brevity but would follow the same pattern)

@login_required
@require_POST
def add_color(request, resume_id):
    resume = get_object_or_404(ResumeTemplate, uuid=resume_id, user=request.user)
    form = TemplateColorForm(request.POST)

    if form.is_valid():
        color = form.save(commit=False)
        color.resume = resume
        color.save()
        messages.success(request, 'Color added successfully!')
    else:
        messages.error(request, 'Error adding color.')

    return redirect('edit_resume', resume_id=resume.uuid)


@login_required
@require_POST
def update_color(request, resume_id, color_id):
    resume = get_object_or_404(ResumeTemplate, uuid=resume_id, user=request.user)
    color = get_object_or_404(TemplateColor, id=color_id, resume=resume)

    form = TemplateColorForm(request.POST, instance=color)

    if form.is_valid():
        form.save()
        messages.success(request, 'Color updated successfully!')
    else:
        messages.error(request, 'Error updating color.')

    return redirect('templates_app:edit_resume', resume_id=resume.uuid)


@login_required
@require_POST
def delete_color(request, resume_id, color_id):
    resume = get_object_or_404(ResumeTemplate, uuid=resume_id, user=request.user)
    color = get_object_or_404(TemplateColor, id=color_id, resume=resume)

    color.delete()
    messages.success(request, 'Color deleted successfully!')

    return redirect('edit_resume', resume_id=resume.uuid)


@login_required
def preview_resume(request, resume_id):
    resume = get_object_or_404(ResumeTemplate, uuid=resume_id, user=request.user)

    return render(request, 'resume_builder/preview_resume.html', {
        'resume': resume
    })


@login_required
def export_resume(request, resume_id):
    resume = get_object_or_404(ResumeTemplate, uuid=resume_id, user=request.user)

    if request.method == 'POST':
        form = ResumeExportForm(request.POST)
        if form.is_valid():
            export_format = form.cleaned_data['export_format']

            # Render the resume HTML
            html_content = render_to_string('resume_builder/export_templates/resume.html', {
                'resume': resume
            })

            if export_format == 'pdf':
                # Generate PDF
                html = HTML(string=html_content)
                css = CSS(string=resume.template.css_styles + resume.custom_css)
                pdf = html.write_pdf(stylesheets=[css])

                # Create response
                response = HttpResponse(pdf, content_type='application/pdf')
                response['Content-Disposition'] = f'attachment; filename="{resume.title}.pdf"'
                return response

            elif export_format == 'docx':
                # Generate DOCX
                doc = Document()
                doc.add_heading(resume.personal_info.full_name, 0)

                # Add personal info
                doc.add_paragraph(resume.personal_info.job_title)
                contact_info = f"{resume.personal_info.email} | {resume.personal_info.phone}"
                doc.add_paragraph(contact_info)

                # Add summary
                if resume.personal_info.summary:
                    doc.add_heading('Summary', level=1)
                    doc.add_paragraph(resume.personal_info.summary)

                # Add experiences
                if resume.experiences.exists():
                    doc.add_heading('Experience', level=1)
                    for exp in resume.experiences.all().order_by('order'):
                        p = doc.add_paragraph()
                        p.add_run(f"{exp.position} at {exp.company}").bold = True
                        p.add_run(f"\n{exp.start_date.strftime('%b %Y')} - ")
                        if exp.current:
                            p.add_run("Present")
                        elif exp.end_date:
                            p.add_run(f"{exp.end_date.strftime('%b %Y')}")
                        if exp.location:
                            p.add_run(f" | {exp.location}")
                        if exp.description:
                            doc.add_paragraph(exp.description)

                # Add education
                if resume.educations.exists():
                    doc.add_heading('Education', level=1)
                    for edu in resume.educations.all().order_by('order'):
                        p = doc.add_paragraph()
                        p.add_run(f"{edu.degree}").bold = True
                        if edu.field_of_study:
                            p.add_run(f" in {edu.field_of_study}")
                        p.add_run(f"\n{edu.institution}")
                        if edu.location:
                            p.add_run(f", {edu.location}")
                        p.add_run(f"\n{edu.start_date.strftime('%b %Y')} - ")
                        if edu.current:
                            p.add_run("Present")
                        elif edu.end_date:
                            p.add_run(f"{edu.end_date.strftime('%b %Y')}")
                        if edu.description:
                            doc.add_paragraph(edu.description)

                # Add skills
                if resume.skills.exists():
                    doc.add_heading('Skills', level=1)
                    skill_list = [skill.name for skill in resume.skills.all().order_by('order')]
                    doc.add_paragraph(', '.join(skill_list))

                # Save to buffer
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                # Create response
                response = HttpResponse(buffer.read(),
                                        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = f'attachment; filename="{resume.title}.docx"'
                return response

            elif export_format == 'txt':
                # Generate plain text
                text_content = []

                # Personal info
                text_content.append(resume.personal_info.full_name.upper())
                text_content.append(resume.personal_info.job_title)
                contact_info = []
                if resume.personal_info.email:
                    contact_info.append(resume.personal_info.email)
                if resume.personal_info.phone:
                    contact_info.append(resume.personal_info.phone)
                if resume.personal_info.address:
                    contact_info.append(resume.personal_info.address)
                text_content.append(' | '.join(contact_info))
                text_content.append('')

                # Summary
                if resume.personal_info.summary:
                    text_content.append('SUMMARY')
                    text_content.append('-' * 80)
                    text_content.append(resume.personal_info.summary)
                    text_content.append('')

                # Experience
                if resume.experiences.exists():
                    text_content.append('EXPERIENCE')
                    text_content.append('-' * 80)
                    for exp in resume.experiences.all().order_by('order'):
                        text_content.append(f"{exp.position} at {exp.company}")
                        date_range = f"{exp.start_date.strftime('%b %Y')} - "
                        if exp.current:
                            date_range += "Present"
                        elif exp.end_date:
                            date_range += f"{exp.end_date.strftime('%b %Y')}"
                        if exp.location:
                            date_range += f" | {exp.location}"
                        text_content.append(date_range)
                        if exp.description:
                            text_content.append(exp.description)
                        text_content.append('')

                # Education
                if resume.educations.exists():
                    text_content.append('EDUCATION')
                    text_content.append('-' * 80)
                    for edu in resume.educations.all().order_by('order'):
                        degree_text = f"{edu.degree}"
                        if edu.field_of_study:
                            degree_text += f" in {edu.field_of_study}"
                        text_content.append(degree_text)

                        institution_text = f"{edu.institution}"
                        if edu.location:
                            institution_text += f", {edu.location}"
                        text_content.append(institution_text)

                        date_range = f"{edu.start_date.strftime('%b %Y')} - "
                        if edu.current:
                            date_range += "Present"
                        elif edu.end_date:
                            date_range += f"{edu.end_date.strftime('%b %Y')}"
                        text_content.append(date_range)

                        if edu.description:
                            text_content.append(edu.description)
                        text_content.append('')

                # Skills
                if resume.skills.exists():
                    text_content.append('SKILLS')
                    text_content.append('-' * 80)
                    skill_list = [skill.name for skill in resume.skills.all().order_by('order')]
                    text_content.append(', '.join(skill_list))
                    text_content.append('')

                # Join all text content
                full_text = '\n'.join(text_content)

                # Create response
                response = HttpResponse(full_text, content_type='text/plain')
                response['Content-Disposition'] = f'attachment; filename="{resume.title}.txt"'
                return response

    # If not POST or form not valid, redirect back to edit page
    return redirect('edit_resume', resume_id=resume.uuid)


@login_required
def reorder_items(request, resume_id):
    resume = get_object_or_404(ResumeTemplate, uuid=resume_id, user=request.user)

    if request.method == 'POST':
        data = json.loads(request.body)
        item_type = data.get('item_type')
        items = data.get('items', [])

        if item_type == 'experience':
            for item in items:
                experience = get_object_or_404(Experience, id=item['id'], resume=resume)
                experience.order = item['order']
                experience.save()
        elif item_type == 'education':
            for item in items:
                education = get_object_or_404(Education, id=item['id'], resume=resume)
                education.order = item['order']
                education.save()
        elif item_type == 'skill':
            for item in items:
                skill = get_object_or_404(Skill, id=item['id'], resume=resume)
                skill.order = item['order']
                skill.save()
        elif item_type == 'project':
            for item in items:
                project = get_object_or_404(Project, id=item['id'], resume=resume)
                project.order = item['order']
                project.save()
        elif item_type == 'language':
            for item in items:
                language = get_object_or_404(Language, id=item['id'], resume=resume)
                language.order = item['order']
                language.save()
        elif item_type == 'certification':
            for item in items:
                certification = get_object_or_404(Certification, id=item['id'], resume=resume)
                certification.order = item['order']
                certification.save()

        return JsonResponse({'status': 'success'})

    return JsonResponse({'status': 'error', 'message': 'Invalid request'}, status=400)


@login_required
def delete_resume(request, resume_id):
    resume = get_object_or_404(ResumeTemplate, uuid=resume_id, user=request.user)

    if request.method == 'POST':
        resume.delete()
        messages.success(request, 'Resume deleted successfully!')
        return redirect('dashboard')

    return render(request, 'resume_builder/delete_resume.html', {'resume': resume})


def template_detail(request, template_slug):
    template = get_object_or_404(ResumeTemplate, slug=template_slug, is_active=True)

    return render(request, 'resume_builder/template_detail.html', {'template': template})

